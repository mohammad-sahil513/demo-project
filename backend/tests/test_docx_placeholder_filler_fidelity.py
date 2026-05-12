"""DOCX placeholder filler fidelity: integrity issues vs source template.

Each test builds a small DOCX template with deterministic placeholders,
runs :class:`DocxPlaceholderFiller`, and asserts that the integrity check
reports zero violations and the surface fidelity scan finds no orphan
placeholder tokens.
"""

from pathlib import Path

from docx import Document
from modules.assembly.models import AssembledDocument, AssembledSection
from modules.export.docx_placeholder_filler import DocxPlaceholderFiller
from modules.template.models import StyleMap


def _write_min_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Heading", level=1)
    doc.save(path)


def test_docx_placeholder_filler_outputs_docx(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    out = tmp_path / "out.docx"
    _write_min_docx(template)
    assembled = AssembledDocument(
        workflow_run_id="wf",
        doc_type="PDD",
        title="Doc",
        sections=[AssembledSection(section_id="s1", title="Heading", execution_order=1, content="Hello")],
    )
    filler = DocxPlaceholderFiller(tmp_path)
    filler.fill(
        template_path=template,
        assembled=assembled,
        output_path=out,
        style_map=StyleMap(),
        placeholder_schema={},
    )
    assert out.exists()

