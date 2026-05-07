"""Phase 9 — assembly + export (offline)."""

from __future__ import annotations

import asyncio
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from core.constants import DocType, TemplateSource
from modules.assembly.assembler import DocumentAssembler
from modules.assembly.models import AssembledDocument
from modules.export.markdown_tables import parse_gfm_table, split_markdown_blocks
from modules.export.renderer import ExportRenderer
from modules.export.types import ExportDocumentInfo, ExportTemplateInfo
from modules.export.xlsx_builder import XlsxBuilder
from modules.template.models import SectionDefinition, StyleMap
from core.exceptions import WorkflowException
from services.workflow_executor import WorkflowExecutor


def test_parse_gfm_table_separator_uses_raw_line() -> None:
    block = "| A | B |\n| --- | --- |\n| x | y |\n"
    assert parse_gfm_table(block) == [["A", "B"], ["x", "y"]]


def test_split_and_parse_markdown_table() -> None:
    md = "Intro line\n\n| A | B |\n| --- | --- |\n| 1 | 2 |\n"
    blocks = split_markdown_blocks(md)
    assert blocks[0][0] == "text"
    assert blocks[1][0] == "table"
    rows = parse_gfm_table(blocks[1][1])
    assert rows == [["A", "B"], ["1", "2"]]


def test_document_assembler_orders_and_emits_warnings_for_missing_generation() -> None:
    plan = [
        SectionDefinition(
            section_id="sec-miss",
            title="Missing",
            execution_order=0,
            output_type="text",
        ),
        SectionDefinition(
            section_id="sec-b",
            title="Second",
            execution_order=2,
            output_type="text",
        ),
        SectionDefinition(
            section_id="sec-a",
            title="First",
            execution_order=1,
            output_type="text",
        ),
    ]
    gen = {
        "sec-a": {
            "output_type": "text",
            "content": "Hello",
            "citations": [{"path": "x", "page": 1}],
        },
        "sec-b": {"output_type": "text", "content": "World"},
    }
    asm = DocumentAssembler()
    outcome = asm.assemble(
        document_filename="BRD.pdf",
        doc_type="PDD",
        section_plan=plan,
        section_generation_results=gen,
    )
    doc = outcome.document
    assert doc.title == "BRD — PDD"
    assert [s.title for s in doc.sections] == ["First", "Second"]
    assert len(outcome.warnings) == 1
    assert outcome.warnings[0].get("code") == "missing_generation"
    assert outcome.warnings[0].get("section_id") == "sec-miss"
    dumped = doc.model_dump()
    assert "citations" not in dumped


def test_export_renderer_docx_inbuilt(tmp_path: Path) -> None:
    from modules.assembly.models import AssembledSection

    storage = tmp_path
    (storage / "outputs").mkdir(parents=True)
    assembled = AssembledDocument(
        title="T",
        doc_type=DocType.PDD.value,
        sections=[
            AssembledSection(
                section_id="s1",
                title="Overview",
                level=1,
                output_type="text",
                content="Body text here.",
            ),
        ],
    )
    renderer = ExportRenderer(storage)
    style = StyleMap()
    out, name, warns = renderer.render(
        workflow_run_id="wf-test",
        document=ExportDocumentInfo(filename="BRD.pdf"),
        template=ExportTemplateInfo(
            template_id="tpl-inbuilt-pdd",
            template_source=TemplateSource.INBUILT,
            file_path=None,
        ),
        assembled=assembled,
        style_map=style,
    )
    assert all(str(w.get("code", "")).startswith("docx_structure") for w in warns)
    assert out.suffix == ".docx"
    assert name.endswith(".docx")
    assert out.exists() and out.stat().st_size > 2000
    loaded = Document(str(out))
    assert any("Overview" in (p.text or "") for p in loaded.paragraphs)


def test_export_renderer_uat_xlsx(tmp_path: Path) -> None:
    storage = tmp_path
    (storage / "outputs").mkdir(parents=True)
    from modules.assembly.models import AssembledSection

    assembled = AssembledDocument(
        title="T",
        doc_type=DocType.UAT.value,
        sections=[
            AssembledSection(
                section_id="u1",
                title="Test Cases",
                level=1,
                output_type="table",
                content="| ID | Name |\n| --- | --- |\n| 1 | Login |\n",
            ),
        ],
    )
    renderer = ExportRenderer(storage)
    out, name, warns = renderer.render(
        workflow_run_id="wf-uat",
        document=ExportDocumentInfo(filename="BRD.pdf"),
        template=ExportTemplateInfo(
            template_id="tpl-inbuilt-uat",
            template_source=TemplateSource.INBUILT,
            file_path=None,
        ),
        assembled=assembled,
        style_map=StyleMap(),
    )
    assert warns == []
    assert out.suffix == ".xlsx"
    assert "UAT" in name
    assert out.exists() and out.stat().st_size > 500


def test_xlsx_builder_emits_schema_warnings(tmp_path: Path) -> None:
    from modules.assembly.models import AssembledSection

    assembled = AssembledDocument(
        title="UAT",
        doc_type=DocType.UAT.value,
        sections=[
            AssembledSection(
                section_id="u1",
                title="Test Cases",
                level=1,
                output_type="table",
                content="| Scenario | Steps |\n| --- | --- |\n| Login | Do login |\n",
            ),
        ],
    )
    out = tmp_path / "warn.xlsx"
    builder = XlsxBuilder()
    warnings = builder.build(
        assembled,
        out,
        sheet_map={
            "schema": [
                {
                    "sheet_name": "Test Cases",
                    "headers": ["ID", "Scenario", "Steps", "Expected Result"],
                    "required_columns": ["ID", "Scenario", "Steps", "Expected Result"],
                }
            ]
        },
    )
    codes = {w.get("code") for w in warnings}
    assert "missing_required_columns" in codes


def test_render_export_blocks_before_output_persistence(tmp_path: Path) -> None:
    out_file = tmp_path / "blocked.xlsx"
    out_file.write_text("temp", encoding="utf-8")

    assembled = AssembledDocument(
        title="UAT",
        doc_type=DocType.UAT.value,
        sections=[
            {
                "section_id": "s1",
                "title": "Test Cases",
                "level": 1,
                "output_type": "table",
                "content": "| A |\n| --- |\n| x |\n",
            }
        ],
    )

    class _WorkflowServiceStub:
        def __init__(self) -> None:
            self.updated = []
            self.wf = SimpleNamespace(
                workflow_run_id="wf-1",
                assembled_document=assembled.model_dump(),
                style_map={},
                warnings=[],
                sheet_map={},
                template_id="tpl-1",
                document_id="doc-1",
                doc_type=DocType.UAT.value,
            )

        def get_or_raise(self, _: str):
            return self.wf

        def get_template(self, _: str):
            return SimpleNamespace(
                template_id="tpl-1",
                template_source=TemplateSource.INBUILT,
                file_path=None,
                placeholder_schema=None,
                resolved_section_bindings=None,
            )

        def get_document(self, _: str):
            return SimpleNamespace(document_id="doc-1", filename="doc.pdf")

        def update(self, workflow_run_id: str, **kwargs):
            self.updated.append((workflow_run_id, kwargs))

    class _OutputServiceStub:
        def __init__(self) -> None:
            self.create_calls = 0

        def create(self, **kwargs):
            self.create_calls += 1
            return SimpleNamespace(output_id="out-1", **kwargs)

    class _RendererStub:
        def render(self, **kwargs):
            return out_file, "blocked.xlsx", [{"code": "missing_required_columns", "severity": "error"}]

    class _EventServiceStub:
        async def emit(self, *_args, **_kwargs):
            return None

    workflow_service = _WorkflowServiceStub()
    output_service = _OutputServiceStub()
    executor = WorkflowExecutor(
        workflow_service=workflow_service,  # type: ignore[arg-type]
        event_service=_EventServiceStub(),  # type: ignore[arg-type]
        output_service=output_service,  # type: ignore[arg-type]
        export_renderer=_RendererStub(),  # type: ignore[arg-type]
    )

    try:
        asyncio.run(executor._phase_render_export("wf-1"))
        raise AssertionError("Expected WorkflowException for blocking export warning")
    except WorkflowException as exc:
        assert "Critical schema mismatch" in str(exc)

    assert output_service.create_calls == 0
    assert not out_file.exists()


def test_render_export_blocks_on_docx_export_policy(tmp_path: Path) -> None:
    out_file = tmp_path / "policy.docx"
    out_file.write_text("temp", encoding="utf-8")

    assembled = AssembledDocument(
        title="PDD",
        doc_type=DocType.PDD.value,
        sections=[
            {
                "section_id": "s1",
                "title": "Overview",
                "level": 1,
                "output_type": "text",
                "content": "Hello",
            }
        ],
    )

    class _WorkflowServiceStub:
        def __init__(self) -> None:
            self.updated = []
            self.wf = SimpleNamespace(
                workflow_run_id="wf-docx-policy",
                assembled_document=assembled.model_dump(),
                style_map={},
                warnings=[],
                sheet_map={},
                template_id="tpl-custom",
                document_id="doc-1",
                doc_type=DocType.PDD.value,
            )

        def get_or_raise(self, _: str):
            return self.wf

        def get_template(self, _: str):
            return SimpleNamespace(
                template_id="tpl-custom",
                template_source=TemplateSource.CUSTOM,
                file_path="x.docx",
                placeholder_schema={},
                resolved_section_bindings={"s1": ["p1"]},
            )

        def get_document(self, _: str):
            return SimpleNamespace(document_id="doc-1", filename="doc.pdf")

        def update(self, workflow_run_id: str, **kwargs):
            self.updated.append((workflow_run_id, kwargs))

    class _OutputServiceStub:
        def __init__(self) -> None:
            self.create_calls = 0

        def create(self, **kwargs):
            self.create_calls += 1
            return SimpleNamespace(output_id="out-1", **kwargs)

    class _RendererStub:
        def render(self, **kwargs):
            return out_file, "policy.docx", [{"code": "docx_legacy_export_disallowed", "phase": "RENDER_EXPORT"}]

    class _EventServiceStub:
        async def emit(self, *_args, **_kwargs):
            return None

    workflow_service = _WorkflowServiceStub()
    output_service = _OutputServiceStub()
    executor = WorkflowExecutor(
        workflow_service=workflow_service,  # type: ignore[arg-type]
        event_service=_EventServiceStub(),  # type: ignore[arg-type]
        output_service=output_service,  # type: ignore[arg-type]
        export_renderer=_RendererStub(),  # type: ignore[arg-type]
    )

    try:
        asyncio.run(executor._phase_render_export("wf-docx-policy"))
        raise AssertionError("Expected WorkflowException for DOCX export policy")
    except WorkflowException as exc:
        assert "template policy" in str(exc).lower()

    assert output_service.create_calls == 0
    assert not out_file.exists()


def test_docx_filler_finds_heading(tmp_path: Path) -> None:
    storage = tmp_path
    tpl_path = tmp_path / "tpl.docx"
    doc = Document()
    doc.add_heading("My Section", level=1)
    doc.add_paragraph("placeholder")
    doc.save(str(tpl_path))

    from modules.assembly.models import AssembledSection
    from modules.export.docx_filler import DocxFiller

    assembled = AssembledDocument(
        title="T",
        doc_type=DocType.PDD.value,
        sections=[
            AssembledSection(
                section_id="s1",
                title="My Section",
                level=1,
                output_type="text",
                content="Replaced body.",
            ),
        ],
    )
    out = tmp_path / "out.docx"
    warns = DocxFiller(storage).fill(tpl_path, assembled, out, style_map=StyleMap())
    assert warns == []
    loaded = Document(str(out))
    texts = [p.text for p in loaded.paragraphs]
    assert any("Replaced body." in t for t in texts)


def test_docx_filler_multiple_sections_re_scans_headings(tmp_path: Path) -> None:
    """Regression: stale paragraph indices must not skip later sections."""
    storage = tmp_path
    tpl_path = tmp_path / "tpl.docx"
    doc = Document()
    doc.add_heading("Alpha", level=1)
    doc.add_paragraph("old alpha body")
    doc.add_heading("Beta", level=1)
    doc.add_paragraph("old beta body")
    doc.save(str(tpl_path))

    from modules.assembly.models import AssembledSection
    from modules.export.docx_filler import DocxFiller

    assembled = AssembledDocument(
        title="T",
        doc_type=DocType.PDD.value,
        sections=[
            AssembledSection(
                section_id="a",
                title="Alpha",
                level=1,
                output_type="text",
                content="New alpha.",
            ),
            AssembledSection(
                section_id="b",
                title="Beta",
                level=1,
                output_type="text",
                content="New beta.",
            ),
        ],
    )
    out = tmp_path / "out.docx"
    warns = DocxFiller(storage).fill(tpl_path, assembled, out, style_map=StyleMap())
    assert warns == []
    loaded = Document(str(out))
    texts = [p.text for p in loaded.paragraphs]
    assert any("New alpha." in t for t in texts)
    assert any("New beta." in t for t in texts)


def test_docx_filler_warns_on_missing_heading(tmp_path: Path) -> None:
    storage = tmp_path
    tpl_path = tmp_path / "tpl.docx"
    doc = Document()
    doc.add_heading("Only This", level=1)
    doc.save(str(tpl_path))

    from modules.assembly.models import AssembledSection
    from modules.export.docx_filler import DocxFiller

    assembled = AssembledDocument(
        title="T",
        doc_type=DocType.PDD.value,
        sections=[
            AssembledSection(
                section_id="x",
                title="Not In Template",
                level=1,
                output_type="text",
                content="orphan",
            ),
        ],
    )
    out = tmp_path / "out.docx"
    warns = DocxFiller(storage).fill(tpl_path, assembled, out)
    assert len(warns) == 1
    assert warns[0].get("code") == "template_heading_not_found"
