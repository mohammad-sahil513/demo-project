"""Tests for DOCX title / TOC / body pagination fixer."""

from __future__ import annotations

import xml.etree.ElementTree as ET
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.text import WD_BREAK

from modules.assembly.models import AssembledDocument, AssembledSection
from modules.export.docx_builder import DocxBuilder
from modules.export.docx_structure_fixer import enforce_title_toc_pagination
from modules.template.models import StyleMap

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{_W}}}{tag}"


def _read_body_children(docx: Path) -> list[ET.Element]:
    with ZipFile(docx, "r") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    body = root.find(f".//{_w('body')}")
    assert body is not None
    return list(body)


def _local(tag: str) -> str:
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _is_page_break_only_p(p: ET.Element) -> bool:
    if _local(p.tag) != "p":
        return False
    runs = [c for c in p if _local(c.tag) == "r"]
    if len(runs) != 1:
        return False
    brs = [c for c in runs[0] if _local(c.tag) == "br"]
    if len(brs) != 1:
        return False
    t = brs[0].get(_w("type")) or brs[0].get("{%s}type" % _W)
    return t == "page"


def _has_toc_instr(p: ET.Element) -> bool:
    for el in p.iter():
        if _local(el.tag) == "instrText" and (el.text or "").strip().upper().startswith("TOC"):
            return True
    return False


def _p_style_val(p: ET.Element) -> str | None:
    ppr = p.find(_w("pPr"))
    if ppr is None:
        return None
    ps = ppr.find(_w("pStyle"))
    if ps is None:
        return None
    return (ps.get(_w("val")) or ps.get("{%s}val" % _W) or "").strip()


def test_inbuilt_build_then_fixer_inserts_page_breaks() -> None:
    assembled = AssembledDocument(
        title="Test Title",
        doc_type="PDD",
        sections=[
            AssembledSection(section_id="s1", title="Section A", level=1, output_type="text", content="Hello."),
        ],
    )
    with TemporaryDirectory() as td:
        out = Path(td) / "out.docx"
        DocxBuilder(Path(td)).build(assembled, StyleMap(), out)
        warns = enforce_title_toc_pagination(out, doc_title="brd.docx", doc_type="PDD")
        codes = {str(w.get("code")) for w in warns}
        assert "docx_structure_pagebreak_inserted" in codes

        kids = _read_body_children(out)
        non_sect = [c for c in kids if _local(c.tag) != "sectPr"]
        assert _p_style_val(non_sect[0]) == "Title"
        assert _is_page_break_only_p(non_sect[1])
        assert _has_toc_instr(non_sect[2])
        assert _is_page_break_only_p(non_sect[3])
        assert _local(non_sect[4].tag) == "p"


def test_fixer_idempotent_second_run_unmodified() -> None:
    assembled = AssembledDocument(
        title="T",
        doc_type="SDD",
        sections=[
            AssembledSection(section_id="s1", title="One", level=1, output_type="text", content="x"),
        ],
    )
    with TemporaryDirectory() as td:
        out = Path(td) / "x.docx"
        DocxBuilder(Path(td)).build(assembled, StyleMap(), out)
        enforce_title_toc_pagination(out, doc_title="f.docx", doc_type="SDD")
        w2 = enforce_title_toc_pagination(out, doc_title="f.docx", doc_type="SDD")
        assert any(str(w.get("code")) == "docx_structure_unmodified" for w in w2)


def test_custom_template_no_toc_gets_toc_and_breaks() -> None:
    with TemporaryDirectory() as td:
        out = Path(td) / "custom.docx"
        doc = Document()
        tp = doc.add_paragraph("Cover")
        tp.style = doc.styles["Title"]
        doc.add_paragraph("First body", style="Heading 1")
        doc.save(out)

        warns = enforce_title_toc_pagination(out, doc_title="custom.docx", doc_type="PDD")
        codes = {str(w.get("code")) for w in warns}
        assert "docx_structure_toc_inserted" in codes or "docx_structure_pagebreak_inserted" in codes

        kids = _read_body_children(out)
        non_sect = [c for c in kids if _local(c.tag) != "sectPr"]
        assert any(_has_toc_instr(p) for p in non_sect if _local(p.tag) == "p")


def test_sectpr_remains_last_body_child() -> None:
    assembled = AssembledDocument(
        title="T",
        doc_type="PDD",
        sections=[AssembledSection(section_id="s1", title="A", level=1, output_type="text", content="b")],
    )
    with TemporaryDirectory() as td:
        out = Path(td) / "s.docx"
        DocxBuilder(Path(td)).build(assembled, StyleMap(), out)
        before = _read_body_children(out)
        assert _local(before[-1].tag) == "sectPr"

        enforce_title_toc_pagination(out, doc_title="s.docx", doc_type="PDD")
        after = _read_body_children(out)
        assert _local(after[-1].tag) == "sectPr"


def test_subtitle_counted_in_title_block_no_duplicate_title_insert() -> None:
    with TemporaryDirectory() as td:
        out = Path(td) / "sub.docx"
        doc = Document()
        t = doc.add_paragraph("Main")
        t.style = doc.styles["Title"]
        try:
            s = doc.add_paragraph("Sub line")
            s.style = doc.styles["Subtitle"]
        except KeyError:
            pytest.skip("Subtitle style not in default template")
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_break(WD_BREAK.PAGE)
        doc.add_paragraph("Body", style="Heading 1")
        doc.save(out)

        warns = enforce_title_toc_pagination(out, doc_title="sub.docx", doc_type="PDD")
        codes = {str(w.get("code")) for w in warns}
        assert "docx_structure_title_inserted" not in codes

        kids = _read_body_children(out)
        non_sect = [c for c in kids if _local(c.tag) != "sectPr"]
        assert _p_style_val(non_sect[0]) == "Title"
        assert _p_style_val(non_sect[1]) == "Subtitle"
        assert _is_page_break_only_p(non_sect[2])


def test_toc_without_trailing_page_break_gets_one() -> None:
    """Title + page + TOC + body (no break after TOC) -> fixer inserts break after TOC."""
    with TemporaryDirectory() as td:
        out = Path(td) / "gap.docx"
        doc = Document()
        tp = doc.add_paragraph("T")
        tp.style = doc.styles["Title"]
        brp = doc.add_paragraph()
        brp.add_run().add_break(WD_BREAK.PAGE)
        # TOC field like builder
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p = doc.add_paragraph()
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = r'TOC \o "1-3" \h \z \u'
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        ft = OxmlElement("w:t")
        ft.text = "TOC"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(ft)
        run._r.append(fld_end)
        doc.add_paragraph("Body text", style="Normal")
        doc.save(out)

        warns = enforce_title_toc_pagination(out, doc_title="gap.docx", doc_type="PDD")
        assert any(str(w.get("code")) == "docx_structure_pagebreak_inserted" for w in warns)

        kids = _read_body_children(out)
        non_sect = [c for c in kids if _local(c.tag) != "sectPr"]
        toc_i = next(i for i, p in enumerate(non_sect) if _local(p.tag) == "p" and _has_toc_instr(p))
        assert _is_page_break_only_p(non_sect[toc_i + 1])
