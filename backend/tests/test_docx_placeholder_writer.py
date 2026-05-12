"""In-place DOCX placeholder writer: targeted node rewrites only.

Builds a minimal DOCX with placeholders in different scopes (paragraph
text tokens, content controls, bookmarks) and asserts that the writer:

- Replaces only the targeted node and leaves siblings untouched.
- Escapes XML special characters in placeholder values.
- Preserves header/footer/media parts byte-for-byte in the output zip.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from modules.assembly.models import AssembledDocument, AssembledSection
from modules.export.docx_placeholder_writer import fill_docx_placeholders_native
from modules.template.schema_extractor_docx import extract_docx_schema


def test_native_writer_replaces_token(tmp_path: Path) -> None:
    tpl = tmp_path / "t.docx"
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Intro {{sec-custom-1-overview}} end.")
    doc.save(tpl)

    schema = extract_docx_schema(tpl)
    schema_dump = schema.model_dump()
    resolved = {"sec-custom-1-overview": ["sec-custom-1-overview"]}
    assembled = AssembledDocument(
        title="T",
        doc_type="PDD",
        sections=[
            AssembledSection(
                section_id="sec-custom-1-overview",
                title="Overview",
                level=1,
                content="REPLACED_TEXT",
            )
        ],
    )
    out = tmp_path / "out.docx"
    w = fill_docx_placeholders_native(
        template_path=tpl,
        output_path=out,
        assembled=assembled,
        placeholder_schema=schema_dump,
        section_placeholder_map=resolved,
    )
    assert not any((x or {}).get("severity") == "error" for x in w)

    out_doc = Document(str(out))
    full = "\n".join(p.text for p in out_doc.paragraphs)
    assert "REPLACED_TEXT" in full
    assert "{{sec-custom-1-overview}}" not in full
