"""Surface-fidelity checks: detect dropped tables or paragraphs vs template.

The renderer must not silently drop content controls or tables that were
present in the original template. These tests build a minimal DOCX with a
table and assert that the surface fidelity check warns when the output
loses that structural element.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from modules.export.docx_integrity import check_docx_surface_fidelity


def test_surface_fidelity_warns_when_output_has_fewer_tables(tmp_path: Path) -> None:
    tpl = tmp_path / "template.docx"
    out = tmp_path / "output.docx"

    d1 = Document()
    d1.add_heading("Section", level=1)
    d1.add_table(rows=2, cols=2)
    d1.save(tpl)

    d2 = Document()
    d2.add_heading("Section", level=1)
    d2.save(out)

    warns = check_docx_surface_fidelity(template_path=tpl, output_path=out, placeholder_schema=None)
    assert any(w.get("code") == "docx_table_count_regression" for w in warns)


def test_surface_fidelity_skipped_when_schema_flag_set(tmp_path: Path) -> None:
    tpl = tmp_path / "template.docx"
    out = tmp_path / "output.docx"
    Document().save(tpl)
    Document().save(out)
    warns = check_docx_surface_fidelity(
        template_path=tpl,
        output_path=out,
        placeholder_schema={"skip_surface_fidelity_checks": True},
    )
    assert warns == []
