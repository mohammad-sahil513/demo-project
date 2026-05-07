"""Custom DOCX export policy gating (legacy / require-native)."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from core.config import settings
from core.constants import DocType, TemplateSource
from modules.assembly.models import AssembledDocument, AssembledSection
from modules.export.renderer import ExportRenderer
from modules.export.types import ExportDocumentInfo, ExportTemplateInfo
from modules.template.models import StyleMap


def _minimal_custom_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Body.")
    doc.save(path)


def test_blocks_legacy_heading_fill_when_disallowed(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(settings, "template_docx_legacy_export_allowed", False)
    monkeypatch.setattr(settings, "template_fidelity_strict_enabled", False)
    monkeypatch.setattr(settings, "template_docx_placeholder_native_enabled", False)
    monkeypatch.setattr(settings, "template_docx_require_native_for_custom", False)

    storage = tmp_path
    tpl_dir = storage / "templates"
    tpl_dir.mkdir(parents=True)
    docx_path = tpl_dir / "custom.docx"
    _minimal_custom_docx(docx_path)

    assembled = AssembledDocument(
        title="T",
        doc_type=DocType.PDD.value,
        sections=[
            AssembledSection(
                section_id="s1",
                title="Overview",
                level=1,
                output_type="text",
                content="Hello",
            ),
        ],
    )
    renderer = ExportRenderer(storage)
    _out, _name, warns = renderer.render(
        workflow_run_id="wf-gate",
        document=ExportDocumentInfo(filename="BRD.pdf"),
        template=ExportTemplateInfo(
            template_id="tpl-custom",
            template_source=TemplateSource.CUSTOM,
            file_path="custom.docx",
            placeholder_schema=None,
            section_placeholder_map=None,
        ),
        assembled=assembled,
        style_map=StyleMap(),
    )
    assert any((w or {}).get("code") == "docx_legacy_export_disallowed" for w in warns)
    assert not _out.exists()


def test_blocks_when_native_required_but_not_available(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(settings, "template_docx_require_native_for_custom", True)
    monkeypatch.setattr(settings, "template_docx_placeholder_native_enabled", True)
    monkeypatch.setattr(settings, "template_fidelity_strict_enabled", False)
    monkeypatch.setattr(settings, "template_docx_legacy_export_allowed", True)

    storage = tmp_path
    tpl_dir = storage / "templates"
    tpl_dir.mkdir(parents=True)
    docx_path = tpl_dir / "custom.docx"
    _minimal_custom_docx(docx_path)

    assembled = AssembledDocument(
        title="T",
        doc_type=DocType.PDD.value,
        sections=[
            AssembledSection(
                section_id="s1",
                title="Overview",
                level=1,
                output_type="text",
                content="Hello",
            ),
        ],
    )
    renderer = ExportRenderer(storage)
    _out, _name, warns = renderer.render(
        workflow_run_id="wf-native",
        document=ExportDocumentInfo(filename="BRD.pdf"),
        template=ExportTemplateInfo(
            template_id="tpl-custom",
            template_source=TemplateSource.CUSTOM,
            file_path="custom.docx",
            placeholder_schema=None,
            section_placeholder_map={},
        ),
        assembled=assembled,
        style_map=StyleMap(),
    )
    assert any((w or {}).get("code") == "docx_native_prerequisites_unmet" for w in warns)
    assert not _out.exists()
