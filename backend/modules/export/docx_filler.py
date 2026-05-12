"""Legacy heading-based DOCX filler for custom templates.

This is the **legacy** path that pre-dates the placeholder-native filler.
It walks the user's DOCX, finds heading paragraphs that match the
:class:`AssembledSection` titles, and writes generated content directly
under each heading.

Used when:

- The template ships without placeholders (no ``{{ tokens }}``, no content
  controls, no named bookmarks), AND
- The legacy export path is enabled (``EXPORT_ALLOW_LEGACY_HEADING_FILL``).

When the strict-fidelity feature flag is on, the renderer refuses to use
this path and forces native filling. The code remains for backward
compatibility with templates uploaded before the placeholder schema
existed.
"""

from __future__ import annotations

import re
import unicodedata
from copy import deepcopy
from pathlib import Path
from dataclasses import asdict

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from core.ids import utc_now_iso
from modules.assembly.models import AssembledDocument
from modules.export.content_blocks import ContentBlock, parse_content_blocks
from modules.export.docx_update_fields import ensure_update_fields_on_open
from modules.template.models import StyleMap

# Fixed diagram display box (requested final-polish behavior)
FIXED_DIAGRAM_WIDTH_IN = 5.8
FIXED_DIAGRAM_HEIGHT_IN = 3.4

# Default non-diagram image width
DEFAULT_IMAGE_WIDTH_IN = 5.5


def _heading_level(paragraph: Paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style else ""
    if not style_name.startswith("Heading"):
        return None
    parts = style_name.split()
    try:
        return int(parts[-1])
    except (ValueError, IndexError):
        return 1


def _is_heading_like(paragraph: Paragraph) -> bool:
    if _heading_level(paragraph) is not None:
        return True

    style_name = paragraph.style.name if paragraph.style else ""
    if style_name and "heading" in style_name.lower():
        return True

    text = paragraph.text.strip()
    if not text:
        return False

    if text.isupper() and len(text.split()) <= 12:
        return True

    return False


def _insert_paragraph_after(paragraph: Paragraph, text: str = "", style: object | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)  # type: ignore[attr-defined]
    new_para = Paragraph(new_p, paragraph._parent)  # type: ignore[arg-type]
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def _clear_following_until_heading(doc: Document, heading: Paragraph) -> None:
    el = heading._p
    nxt = el.getnext()
    while nxt is not None:
        if nxt.tag.endswith("p"):
            para = Paragraph(nxt, doc)
            if _heading_level(para) is not None:
                break
        nxt_next = nxt.getnext()
        parent = nxt.getparent()
        if parent is not None:
            parent.remove(nxt)
        nxt = nxt_next


def _normalize_heading_text(value: str) -> str:
    text = unicodedata.normalize("NFKC", value or "").strip().lower()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(
        r"^(?:[a-z]\.|(?:\d+(?:\.\d+)*)[.)]?)\s+",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"^(?:\\)?#{1,6}\s+", "", text)
    text = re.sub(r"[^\w\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _heading_match_score(template_heading: str, section_title: str) -> int:
    raw_template = (template_heading or "").strip()
    raw_target = (section_title or "").strip()
    if not raw_template or not raw_target:
        return -1

    if raw_template == raw_target:
        return 100

    norm_template = _normalize_heading_text(raw_template)
    norm_target = _normalize_heading_text(raw_target)
    if not norm_template or not norm_target:
        return -1

    if norm_template == norm_target:
        return 90

    if norm_template in norm_target or norm_target in norm_template:
        shorter = min(len(norm_template), len(norm_target))
        if shorter >= 6:
            return 70

    return -1


def _find_best_heading_paragraph(
    doc: Document,
    section_title: str,
    used_heading_elements: set[int],
) -> Paragraph | None:
    best_para: Paragraph | None = None
    best_score = -1

    for p in doc.paragraphs:
        if id(p._p) in used_heading_elements:
            continue

        text = p.text.strip()
        if not text:
            continue

        score = _heading_match_score(text, section_title)
        if score < 0:
            continue

        if _is_heading_like(p):
            score += 5

        if score > best_score:
            best_score = score
            best_para = p
            if best_score >= 100:
                break

    return best_para


def _apply_run_style(
    run,
    *,
    font_name: str,
    font_size_pt: int,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.italic = italic


def _apply_paragraph_alignment(paragraph: Paragraph, alignment: str) -> None:
    value = (alignment or "").strip().lower()
    if value == "center":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif value == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def _apply_body_style(paragraph: Paragraph, sm: StyleMap, *, relax_emphasis: bool = False) -> None:
    bold = sm.body.bold and not relax_emphasis
    italic = sm.body.italic and not relax_emphasis
    for run in paragraph.runs:
        _apply_run_style(
            run,
            font_name=sm.body.font_name,
            font_size_pt=sm.body.font_size_pt,
            bold=bold,
            italic=italic,
        )
    _apply_paragraph_alignment(paragraph, sm.body.alignment)


def _apply_heading_style(paragraph: Paragraph, sm: StyleMap, level: int) -> None:
    style_cfg = sm.heading_1 if level <= 1 else sm.heading_2
    for run in paragraph.runs:
        _apply_run_style(
            run,
            font_name=style_cfg.font_name,
            font_size_pt=style_cfg.font_size_pt,
            bold=style_cfg.bold,
            italic=style_cfg.italic,
        )
    _apply_paragraph_alignment(paragraph, style_cfg.alignment)


def _apply_caption_style(paragraph: Paragraph, sm: StyleMap) -> None:
    for run in paragraph.runs:
        _apply_run_style(
            run,
            font_name=sm.body.font_name,
            font_size_pt=max(sm.body.font_size_pt - 1, 9),
            bold=False,
            italic=True,
        )
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _add_line_after_doc(
    doc: Document,
    after: Paragraph,
    line: str,
    normal,
    sm: StyleMap,
    *,
    relax_emphasis: bool = False,
) -> Paragraph:
    stripped = line.strip()
    if not stripped:
        return after

    bullet_prefixes = ("- ", "* ", "• ")
    numbered = re.match(r"^(\d+[.)])\s+", stripped)

    ref = _insert_paragraph_after(after, "", normal)

    for bp in bullet_prefixes:
        if stripped.startswith(bp):
            text = stripped[len(bp):].strip()
            try:
                ref.style = doc.styles["List Bullet"]
                if text:
                    ref.add_run(text)
            except (KeyError, ValueError):
                ref.add_run(f"• {text}")
            _apply_body_style(ref, sm, relax_emphasis=relax_emphasis)
            return ref

    if numbered:
        text = stripped[numbered.end():].strip()
        try:
            ref.style = doc.styles["List Number"]
            if text:
                ref.add_run(text)
        except (KeyError, ValueError):
            ref.add_run(stripped)
        _apply_body_style(ref, sm, relax_emphasis=relax_emphasis)
        return ref

    out = _insert_paragraph_after(after, stripped, normal)
    _apply_body_style(out, sm, relax_emphasis=relax_emphasis)
    return out


def _build_table_element(rows: list[list[str]], sm: StyleMap):
    scratch = Document()
    cols = max(len(r) for r in rows)
    table = scratch.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            val = row[c_idx] if c_idx < len(row) else ""
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(sm.body.font_size_pt)
                    run.font.name = sm.body.font_name
                    if r_idx == 0:
                        run.font.bold = True

    return deepcopy(table._tbl)


def _normalize_image_target_key(value: str | None) -> str | None:
    if not value:
        return None
    return str(Path(value)).replace("\\", "/").strip().lower()


class DocxFiller:
    def __init__(self, storage_root: Path) -> None:
        self._storage_root = storage_root
        self._figure_counter = 0

    def fill(
        self,
        template_path: Path,
        assembled: AssembledDocument,
        output_path: Path,
        *,
        style_map: StyleMap | None = None,
    ) -> list[dict[str, object]]:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document(str(template_path))
        used_heading_elements: set[int] = set()
        warnings: list[dict[str, object]] = []
        sm = style_map or StyleMap()
        ts = utc_now_iso()
        normal = doc.styles["Normal"]

        for section in assembled.sections:
            heading = _find_best_heading_paragraph(doc, section.title, used_heading_elements)
            if heading is None:
                warnings.append(
                    {
                        "phase": "RENDER_EXPORT",
                        "code": "template_heading_not_found",
                        "section_id": section.section_id,
                        "title": section.title,
                        "message": (
                            f"No paragraph matching heading {section.title!r} in template; "
                            "section content was not inserted."
                        ),
                        "at": ts,
                    },
                )
                continue

            used_heading_elements.add(id(heading._p))
            _clear_following_until_heading(doc, heading)
            ref: Paragraph = heading

            diagram_rendered = False
            if section.output_type == "diagram" and section.diagram_path:
                img = self._storage_root / Path(section.diagram_path)
                if img.exists():
                    ref = self._insert_image_with_caption(
                        doc=doc,
                        after=ref,
                        image_path=img,
                        caption=f"Figure {self._next_figure_number()}: {section.title}",
                        sm=sm,
                        normal=normal,
                        fixed_diagram_size=True,
                    )
                    diagram_rendered = True

            blocks = section.content_blocks or [asdict(b) for b in parse_content_blocks(section.content)]
            ref = self._append_blocks(
                doc=doc,
                after=ref,
                section_title=section.title,
                content_blocks=blocks,
                sm=sm,
                normal=normal,
                suppress_image_target=_normalize_image_target_key(section.diagram_path) if diagram_rendered else None,
                suppress_all_image_blocks=diagram_rendered and section.output_type == "diagram",
                diagram_section=(section.output_type == "diagram"),
                relax_emphasis=True,
            )

        doc.save(str(output_path))
        ensure_update_fields_on_open(output_path)
        return warnings

    def _append_blocks(
        self,
        doc: Document,
        after: Paragraph,
        *,
        section_title: str,
        content_blocks: list[dict[str, object]],
        sm: StyleMap,
        normal,
        suppress_image_target: str | None = None,
        suppress_all_image_blocks: bool = False,
        diagram_section: bool = False,
        relax_emphasis: bool = True,
    ) -> Paragraph:
        ref = after

        for raw_block in content_blocks:
            block = self._coerce_block(raw_block)
            if block is None:
                continue

            if block.kind == "heading":
                if _normalize_heading_text(block.text) == _normalize_heading_text(section_title):
                    continue
                ref = self._insert_heading_after(doc, ref, block.text, block.level, sm)
                continue

            if block.kind == "paragraph":
                para = _insert_paragraph_after(ref, block.text, normal)
                _apply_body_style(para, sm, relax_emphasis=relax_emphasis)
                ref = para
                continue

            if block.kind == "bullet_list":
                for item in block.items:
                    ref = _add_line_after_doc(
                        doc, ref, f"- {item}", normal, sm, relax_emphasis=relax_emphasis
                    )
                continue

            if block.kind == "numbered_list":
                for idx, item in enumerate(block.items, start=1):
                    ref = _add_line_after_doc(
                        doc, ref, f"{idx}. {item}", normal, sm, relax_emphasis=relax_emphasis
                    )
                continue

            if block.kind in {"table_gfm", "table_html"}:
                if not block.rows:
                    continue
                tbl_el = _build_table_element(block.rows, sm)
                ref._p.addnext(tbl_el)  # type: ignore[attr-defined]

                tail = OxmlElement("w:p")
                tbl_el.addnext(tail)
                ref = Paragraph(tail, doc)
                continue

            if block.kind == "image":
                if suppress_all_image_blocks:
                    continue

                target_key = _normalize_image_target_key(block.image_target)
                if suppress_image_target and target_key == suppress_image_target:
                    continue

                image_path = self._resolve_image_path(block.image_target)
                if image_path is None or not image_path.exists():
                    continue

                alt = block.image_alt or section_title
                ref = self._insert_image_with_caption(
                    doc=doc,
                    after=ref,
                    image_path=image_path,
                    caption=f"Figure {self._next_figure_number()}: {alt}",
                    sm=sm,
                    normal=normal,
                    fixed_diagram_size=diagram_section,
                )
                continue

            if block.kind == "caption":
                continue

        return ref

    def _insert_heading_after(
        self,
        doc: Document,
        after: Paragraph,
        text: str,
        level: int,
        sm: StyleMap,
    ) -> Paragraph:
        para = _insert_paragraph_after(after, text, None)
        style_name = "Heading 2" if level <= 2 else "Heading 3"
        try:
            para.style = doc.styles[style_name]
        except (KeyError, ValueError):
            para.style = doc.styles["Normal"]
        _apply_heading_style(para, sm, level=1 if level <= 2 else 2)
        return para

    def _insert_image_with_caption(
        self,
        *,
        doc: Document,
        after: Paragraph,
        image_path: Path,
        caption: str,
        sm: StyleMap,
        normal,
        fixed_diagram_size: bool = False,
    ) -> Paragraph:
        img_para = _insert_paragraph_after(after, "", normal)
        run = img_para.add_run()

        if fixed_diagram_size:
            run.add_picture(
                str(image_path),
                width=Inches(FIXED_DIAGRAM_WIDTH_IN),
                height=Inches(FIXED_DIAGRAM_HEIGHT_IN),
            )
        else:
            run.add_picture(
                str(image_path),
                width=Inches(DEFAULT_IMAGE_WIDTH_IN),
            )

        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        caption_para = _insert_paragraph_after(img_para, caption, None)
        try:
            caption_para.style = doc.styles["Caption"]
        except (KeyError, ValueError):
            caption_para.style = doc.styles["Normal"]
        _apply_caption_style(caption_para, sm)
        return caption_para

    def _next_figure_number(self) -> int:
        self._figure_counter += 1
        return self._figure_counter

    def _resolve_image_path(self, image_target: str | None) -> Path | None:
        if not image_target:
            return None

        target = Path(image_target)
        if target.is_absolute():
            return target

        direct = self._storage_root / target
        if direct.exists():
            return direct

        candidate = self._storage_root / "outputs" / target
        if candidate.exists():
            return candidate

        candidate = self._storage_root / "diagrams" / target
        if candidate.exists():
            return candidate

        return None

    def _coerce_block(self, raw_block) -> ContentBlock | None:
        try:
            if isinstance(raw_block, ContentBlock):
                return raw_block

            if not isinstance(raw_block, dict):
                return None

            kind = str(raw_block.get("kind") or "").strip()
            if not kind:
                return None

            text = str(raw_block.get("text") or "")
            level = int(raw_block.get("level") or 0)

            raw_items = raw_block.get("items") or []
            items = [str(item) for item in raw_items] if isinstance(raw_items, list) else []

            raw_rows = raw_block.get("rows") or []
            rows: list[list[str]] = []
            if isinstance(raw_rows, list):
                for row in raw_rows:
                    if isinstance(row, list):
                        rows.append([str(cell) for cell in row])

            image_target = raw_block.get("image_target")
            image_alt = raw_block.get("image_alt")

            return ContentBlock(
                kind=kind,  # type: ignore[arg-type]
                text=text,
                level=level,
                items=items,
                rows=rows,
                image_target=str(image_target) if image_target is not None else None,
                image_alt=str(image_alt) if image_alt is not None else None,
            )
        except Exception:
            return None
