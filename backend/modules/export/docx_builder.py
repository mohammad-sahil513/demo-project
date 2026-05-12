"""Build a fresh DOCX from a StyleMap + assembled sections (inbuilt templates).

This builder runs for inbuilt PDD / SDD / UAT exports where there is no
user-supplied template file. The :class:`StyleMap` registered under
``modules.template.inbuilt.styles`` controls page setup, heading styles,
and table appearance.

The builder consumes :class:`AssembledSection` objects in order. Section
content is parsed into structured :class:`ContentBlock` items (paragraphs,
bullets, tables, images) before being emitted to the document — that keeps
the writer dumb and easy to test.

Tables come in as GFM markdown and are converted to native Word tables.
Diagrams come in as PNG paths; the builder embeds them at a fixed display
box so the visual rhythm of the document stays consistent across sections.
"""
from __future__ import annotations

from dataclasses import asdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from modules.assembly.models import AssembledDocument
from modules.export.content_blocks import ContentBlock, parse_content_blocks
from modules.template.models import StyleMap


# Fixed diagram display box (requested final-polish behavior)
FIXED_DIAGRAM_WIDTH_IN = 5.8
FIXED_DIAGRAM_HEIGHT_IN = 3.4

# Default non-diagram image width
DEFAULT_IMAGE_WIDTH_IN = 5.5


def _apply_run_style(
    run,
    *,
    font_name: str,
    font_size_pt: int,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.italic = italic


def _apply_paragraph_alignment(paragraph, alignment: str) -> None:
    value = (alignment or "").strip().lower()
    if value == "center":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif value == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def _apply_body_style(paragraph, sm: StyleMap) -> None:
    for run in paragraph.runs:
        _apply_run_style(
            run,
            font_name=sm.body.font_name,
            font_size_pt=sm.body.font_size_pt,
            bold=sm.body.bold,
            italic=sm.body.italic,
        )
    _apply_paragraph_alignment(paragraph, sm.body.alignment)


def _apply_heading_style(paragraph, sm: StyleMap, level: int) -> None:
    style_cfg = sm.heading_1 if level <= 1 else sm.heading_2
    for run in paragraph.runs:
        _apply_run_style(
            run,
            font_name=style_cfg.font_name,
            font_size_pt=style_cfg.font_size_pt,
            bold=style_cfg.bold,
            italic=style_cfg.italic,
        )
    _apply_paragraph_alignment(paragraph, style_cfg.alignment)


def _apply_caption_style(paragraph, sm: StyleMap) -> None:
    for run in paragraph.runs:
        _apply_run_style(
            run,
            font_name=sm.body.font_name,
            font_size_pt=max(sm.body.font_size_pt - 1, 9),
            bold=False,
            italic=True,
        )
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _heading_style_name(level: int) -> str:
    if level <= 1:
        return "Heading 1"
    if level == 2:
        return "Heading 2"
    if level == 3:
        return "Heading 3"
    return "Heading 4"


def _normalize_heading_text(value: str) -> str:
    import re
    import unicodedata

    text = unicodedata.normalize("NFKC", value or "").strip().lower()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"^(?:\\)?#{1,6}\s+", "", text)
    text = re.sub(
        r"^(?:[a-z]\.|(?:\d+(?:\.\d+)*)[.)]?)\s+",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"[^\w\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _normalize_image_target_key(value: str | None) -> str | None:
    if not value:
        return None
    return str(Path(value)).replace("\\", "/").strip().lower()


def _add_toc(doc: Document) -> None:
    """
    Insert a Word TOC field.
    Word updates the TOC on open / field update.
    """
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "Table of Contents"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_text)
    run._r.append(fld_end)


def _apply_page_setup(doc: Document, style_map: StyleMap) -> None:
    ps = getattr(style_map, "page_setup", None)
    if ps is None:
        return

    section = doc.sections[0]

    for attr_name, target_attr in (
        ("margin_left_in", "left_margin"),
        ("margin_right_in", "right_margin"),
        ("margin_top_in", "top_margin"),
        ("margin_bottom_in", "bottom_margin"),
    ):
        value = getattr(ps, attr_name, None)
        if value is not None:
            setattr(section, target_attr, Inches(float(value)))

    orientation = getattr(ps, "orientation", None)
    if isinstance(orientation, str) and orientation.strip().lower() == "landscape":
        from docx.enum.section import WD_ORIENT

        section.orientation = WD_ORIENT.LANDSCAPE
        width = section.page_height
        height = section.page_width
        section.page_width = width
        section.page_height = height


class DocxBuilder:
    def __init__(self, storage_root: Path) -> None:
        self._storage_root = storage_root
        self._figure_counter = 0

    def build(self, assembled: AssembledDocument, style_map: StyleMap, output_path: Path) -> None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        _apply_page_setup(doc, style_map)

        # Title
        title_para = doc.add_paragraph(assembled.title)
        try:
            title_para.style = doc.styles["Title"]
        except (KeyError, ValueError):
            pass

        for run in title_para.runs:
            _apply_run_style(
                run,
                font_name=style_map.heading_1.font_name,
                font_size_pt=max(style_map.heading_1.font_size_pt + 2, 18),
                bold=True,
                italic=False,
            )
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # TOC
        _add_toc(doc)

        # Content
        for idx, section in enumerate(assembled.sections):
            heading = doc.add_paragraph(section.title)
            try:
                heading.style = doc.styles[_heading_style_name(section.level)]
            except (KeyError, ValueError):
                try:
                    heading.style = doc.styles["Heading 1"]
                except (KeyError, ValueError):
                    pass
            _apply_heading_style(heading, style_map, level=section.level)

            diagram_rendered = False
            if section.output_type == "diagram" and section.diagram_path:
                img = self._resolve_image_path(section.diagram_path)
                if img and img.exists():
                    self._insert_image_with_caption(
                        doc=doc,
                        image_path=img,
                        caption=f"Figure {self._next_figure_number()}: {section.title}",
                        style_map=style_map,
                        fixed_diagram_size=True,
                    )
                    diagram_rendered = True

            blocks = section.content_blocks or [asdict(b) for b in parse_content_blocks(section.content)]
            self._add_body_blocks(
                doc=doc,
                section_title=section.title,
                blocks=blocks,
                style_map=style_map,
                suppress_image_target=_normalize_image_target_key(section.diagram_path) if diagram_rendered else None,
                suppress_all_image_blocks=diagram_rendered and section.output_type == "diagram",
                diagram_section=(section.output_type == "diagram"),
            )

        doc.save(str(output_path))

    def _add_body_blocks(
        self,
        doc: Document,
        *,
        section_title: str,
        blocks: list[dict[str, object]],
        style_map: StyleMap,
        suppress_image_target: str | None = None,
        suppress_all_image_blocks: bool = False,
        diagram_section: bool = False,
    ) -> None:
        for raw_block in blocks:
            block = self._coerce_block(raw_block)
            if block is None:
                continue

            if block.kind == "heading":
                if _normalize_heading_text(block.text) == _normalize_heading_text(section_title):
                    continue
                para = doc.add_paragraph(block.text)
                try:
                    para.style = doc.styles[_heading_style_name(max(block.level, 2))]
                except (KeyError, ValueError):
                    try:
                        para.style = doc.styles["Heading 2"]
                    except (KeyError, ValueError):
                        pass
                _apply_heading_style(para, style_map, level=2 if block.level <= 2 else 3)
                continue

            if block.kind == "paragraph":
                para = doc.add_paragraph(block.text)
                try:
                    para.style = doc.styles["Normal"]
                except (KeyError, ValueError):
                    pass
                _apply_body_style(para, style_map)
                continue

            if block.kind == "bullet_list":
                for item in block.items:
                    try:
                        para = doc.add_paragraph(item, style="List Bullet")
                    except (KeyError, ValueError):
                        para = doc.add_paragraph(f"• {item}")
                    _apply_body_style(para, style_map)
                continue

            if block.kind == "numbered_list":
                for index, item in enumerate(block.items, start=1):
                    try:
                        para = doc.add_paragraph(item, style="List Number")
                    except (KeyError, ValueError):
                        para = doc.add_paragraph(f"{index}. {item}")
                    _apply_body_style(para, style_map)
                continue

            if block.kind in {"table_gfm", "table_html"}:
                if not block.rows:
                    continue

                cols = max(len(row) for row in block.rows)
                if cols <= 0:
                    continue

                table = doc.add_table(rows=len(block.rows), cols=cols)
                table.style = "Table Grid"

                for r_idx, row in enumerate(block.rows):
                    for c_idx in range(cols):
                        cell = table.cell(r_idx, c_idx)
                        value = row[c_idx] if c_idx < len(row) else ""
                        cell.text = str(value)

                        for para in cell.paragraphs:
                            for run in para.runs:
                                _apply_run_style(
                                    run,
                                    font_name=style_map.body.font_name,
                                    font_size_pt=style_map.body.font_size_pt,
                                    bold=(r_idx == 0),
                                    italic=False,
                                )
                continue

            if block.kind == "image":
                if suppress_all_image_blocks:
                    continue

                target_key = _normalize_image_target_key(block.image_target)
                if suppress_image_target and target_key == suppress_image_target:
                    continue

                image_path = self._resolve_image_path(block.image_target)
                if image_path and image_path.exists():
                    alt = block.image_alt or section_title
                    self._insert_image_with_caption(
                        doc=doc,
                        image_path=image_path,
                        caption=f"Figure {self._next_figure_number()}: {alt}",
                        style_map=style_map,
                        fixed_diagram_size=diagram_section,
                    )
                continue

            if block.kind == "caption":
                continue

    def _insert_image_with_caption(
        self,
        *,
        doc: Document,
        image_path: Path,
        caption: str,
        style_map: StyleMap,
        fixed_diagram_size: bool = False,
    ) -> None:
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = pic_para.add_run()

        if fixed_diagram_size:
            run.add_picture(
                str(image_path),
                width=Inches(FIXED_DIAGRAM_WIDTH_IN),
                height=Inches(FIXED_DIAGRAM_HEIGHT_IN),
            )
        else:
            run.add_picture(
                str(image_path),
                width=Inches(DEFAULT_IMAGE_WIDTH_IN),
            )

        caption_para = doc.add_paragraph(caption)
        try:
            caption_para.style = doc.styles["Caption"]
        except (KeyError, ValueError):
            try:
                caption_para.style = doc.styles["Normal"]
            except (KeyError, ValueError):
                pass
        _apply_caption_style(caption_para, style_map)

    def _resolve_image_path(self, image_target: str | None) -> Path | None:
        if not image_target:
            return None

        target = Path(image_target)
        if target.is_absolute():
            return target

        direct = self._storage_root / target
        if direct.exists():
            return direct

        candidate = self._storage_root / "outputs" / target
        if candidate.exists():
            return candidate

        candidate = self._storage_root / "diagrams" / target
        if candidate.exists():
            return candidate

        return None

    def _next_figure_number(self) -> int:
        self._figure_counter += 1
        return self._figure_counter

    def _coerce_block(self, raw_block: dict[str, object]) -> ContentBlock | None:
        try:
            kind = str(raw_block.get("kind") or "").strip()
            if not kind:
                return None

            return ContentBlock(
                kind=kind,  # type: ignore[arg-type]
                text=str(raw_block.get("text") or ""),
                level=int(raw_block.get("level") or 0),
                items=[str(v) for v in (raw_block.get("items") or [])]
                if isinstance(raw_block.get("items"), list)
                else [],
                rows=[
                    [str(cell) for cell in row]
                    for row in (raw_block.get("rows") or [])
                ]
                if isinstance(raw_block.get("rows"), list)
                else [],
                image_target=str(raw_block.get("image_target"))
                if raw_block.get("image_target") is not None
                else None,
                image_alt=str(raw_block.get("image_alt"))
                if raw_block.get("image_alt") is not None
                else None,
            )
        except Exception:
            return None
